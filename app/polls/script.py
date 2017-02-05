from docx import Document
import re

def open_docx(filename):
	doc = Document(filename)
	return doc

def get_fields(filename):
	doc = open_docx(filename)
	fields = []
	for p in doc.paragraphs:
		tags = re.findall('\[\[[a-zA-Z0-9]*\]\]', p.text)
		for t in tags:
			fields.append(t[2:-2])
	return fields

def replace(text, mapping):
	mod = []
	i = 0
	l = len(text)
	while i < l:
		if text[i] == '[' and i+1<len(text) and text[i+1] == '[':
			i += 2
			var = []
			while not (text[i]==']' and text[i+1]==']'):
				var.append(text[i])
				i+=1
			i+=2
			var = ''.join(var)
			mod += list(mapping[var])
		else:
			mod.append(text[i])
			i+=1
	return ''.join(mod)

def create_file(filename, mapping, id):
	doc = open_docx(filename)
	for p in doc.paragraphs:
		p.text = replace(p.text, mapping)
	n = filename[:-5]+ '_filled'+str(id)+'.docx'
	doc.save(n)
	return n[6:]

def get_mappings(data_filename, vars):
	lines = None
	mappings = []
	with open(data_filename, 'r') as f:
		lines = f.readlines()
		for line in lines:
			mapping = {}
			for token,v in zip(line.split(','), vars):
				mapping[v] = token.strip()
			mappings.append(mapping)
	print(mappings)
	return mappings

def go(vnames):
	f = 'polls/demo.docx'
	#vars = ['name', 'violation', 'pronoun']
	mappings = get_mappings('polls/data.csv', vnames)
	flist = []
	for i in range(len(mappings)):
		flist.append(create_file(f, mappings[i], i))
	return flist


